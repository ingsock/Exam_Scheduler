from typing import Dict, List
from scheduler import Schedule

try:
    from docx import Document

    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx not installed. Install with: pip install python-docx")
    Document = None


from datetime import datetime, timedelta


def generate_word_schedule(
    schedule: Schedule,
    start_date_str: str = "2025-01-01",
    filename: str = "Exam_Schedule.docx",
):
    if not Document:
        print("Cannot generate Word document: python-docx library missing.")
        return

    doc = Document()

    try:
        start_date = datetime.strptime(start_date_str, "%Y-%m-%d")
    except ValueError:
        start_date = datetime.now()

    # Title
    title = doc.add_heading("Exam Schedule", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sort assignments by day
    schedule_data = []  # (Day, Subject Name)
    for subj, day in schedule.assignments.items():
        schedule_data.append((day, subj.name))

    schedule_data.sort(key=lambda x: x[0])

    # Create Table
    # Columns: Date, Subjects
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Date"
    hdr_cells[1].text = "Subjects"

    # Group by Day
    grouped: Dict[int, List[str]] = {}
    for day, subj in schedule_data:
        if day not in grouped:
            grouped[day] = []
        grouped[day].append(subj)

    sorted_days = sorted(grouped.keys())

    for day_idx in sorted_days:
        row_cells = table.add_row().cells

        # Calculate Date
        current_date = start_date + timedelta(days=day_idx)
        date_str = current_date.strftime("%a %d/%m")

        row_cells[0].text = date_str
        row_cells[1].text = ", ".join(grouped[day_idx])

    doc.save(filename)
    return filename
